
import PyPDF2
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ChatMessageHistory, ConversationBufferMemory
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from docx import Document
import pandas as pd
import io
import logging
import os
import requests

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize HuggingFace SentenceTransformer for embeddings
embeddings = HuggingFaceEmbeddings(model_name='all-MiniLM-L6-v2')

# Global variable to store the vector store
global_vector_store = None

groq_api_key = os.environ['GROQ_API_KEY']

# Initializing GROQ chat with provided API key, model name, and settings
llm_groq = ChatGroq(
    groq_api_key=groq_api_key,
    model_name="llama3-70b-8192",
    temperature=0.2
)

def process_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def process_docx(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def process_csv(file_path):
    df = pd.read_csv(file_path)
    buffer = io.StringIO()
    df.info(buf=buffer)
    text = buffer.getvalue()
    text += "\n\nData Summary:\n" + df.describe().to_string()
    return text

def process_file(file_path):
    try:
        if file_path.lower().endswith('.pdf'):
            return process_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return process_docx(file_path)
        elif file_path.lower().endswith('.csv'):
            return process_csv(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path}")
    except Exception as e:
        logger.error(f"Error processing file {file_path}: {str(e)}")
        raise

def process_files(file_paths):
    texts = []
    metadatas = []
    for file_path in file_paths:
        text = process_file(file_path)
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=50)
        file_texts = text_splitter.split_text(text)
        texts.extend(file_texts)
        file_metadatas = [{"source": f"{i}-{os.path.basename(file_path)}"} for i in range(len(file_texts))]
        metadatas.extend(file_metadatas)
    return texts, metadatas

def create_vectorstore(texts, metadatas):
    return FAISS.from_texts(texts, embeddings, metadatas=metadatas)

def process_pdf_content(content):
    pdf_file = io.BytesIO(content)
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def process_docx_content(content):
    docx_file = io.BytesIO(content)
    doc = Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def process_csv_content(content):
    csv_file = io.StringIO(content.decode('utf-8'))
    df = pd.read_csv(csv_file)
    buffer = io.StringIO()
    df.info(buf=buffer)
    text = buffer.getvalue()
    text += "\n\nData Summary:\n" + df.describe().to_string()
    return text

def store_files_api(file_paths, s3_file_urls):
    global global_vector_store
    try:
        logger.info(f"Processing {len(file_paths)} local files and {len(s3_file_urls)} S3 files")
        texts = []
        metadatas = []

        # Process local files
        if file_paths:
            local_texts, local_metadatas = process_files(file_paths)
            texts.extend(local_texts)
            metadatas.extend(local_metadatas)

        # Process S3 file URLs
        for url in s3_file_urls:
            try:
                response = requests.get(url)
                response.raise_for_status()
                content = response.content
                file_name = url.split('/')[-1]

                if url.lower().endswith('.pdf'):
                    text = process_pdf_content(content)
                elif url.lower().endswith('.docx'):
                    text = process_docx_content(content)
                elif url.lower().endswith('.csv'):
                    text = process_csv_content(content)
                else:
                    logger.warning(f"Unsupported file type: {url}")
                    continue

                text_splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=50)
                file_texts = text_splitter.split_text(text)
                texts.extend(file_texts)
                file_metadatas = [{"source": f"{i}-{file_name}"} for i in range(len(file_texts))]
                metadatas.extend(file_metadatas)

            except requests.RequestException as e:
                logger.error(f"Error downloading file from {url}: {str(e)}")

        if not texts:
            raise ValueError("No valid texts extracted from files")

        global_vector_store = create_vectorstore(texts, metadatas)
        logger.info("Vector store created successfully")
        
        return f"Processing complete. {len(texts)} text chunks added to the RAG system. You can now ask questions!"
    except Exception as e:
        logger.error(f"Error in store_files_api: {str(e)}")
        raise

def setup_qa_chain(docsearch):
    message_history = ChatMessageHistory()
    memory = ConversationBufferMemory(
        memory_key="chat_history",
        output_key="answer",
        chat_memory=message_history,
        return_messages=True,
    )
    chain = ConversationalRetrievalChain.from_llm(
        llm=llm_groq,
        chain_type="stuff",
        retriever=docsearch.as_retriever(),
        memory=memory,
        return_source_documents=True,
    )
    return chain

def file_response_api(query):
    global global_vector_store
    
    if global_vector_store is None:
        return "No file has been uploaded yet. Please upload file before asking questions."
    
    qa_chain = setup_qa_chain(global_vector_store)
    res = qa_chain({"question": query})
    answer = res["answer"]
    source_documents = res["source_documents"]
    
    print("\nAnswer:", answer)
    
    if source_documents:
        print("\nSources:")
        for idx, doc in enumerate(source_documents):
            print(f"Source {idx + 1}: {doc.metadata['source']}")
    else:
        print("\nNo sources found")
    
    print("\n" + "-"*50 + "\n")
    
    return answer

def reset_system():
    global global_vector_store
    global_vector_store = None
    logger.info("System reset successfully")